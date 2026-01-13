from pathlib import Path

from docx import Document
from docx.shared import Pt


def add_paragraphs(doc: Document, text: str) -> None:
    for line in text.split("\n"):
        doc.add_paragraph(line)


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    docs_dir = root / "docs"
    docs_dir.mkdir(exist_ok=True)
    file_path = docs_dir / "Venue_Booking_System_SRS.docx"

    doc = Document()

    for level, size in [(1, 18), (2, 16), (3, 14)]:
        style = doc.styles[f"Heading {level}"]
        style.font.size = Pt(size)

    doc.add_heading("Software Requirements Specification (SRS)", 0)
    doc.add_paragraph("Venue Booking System")

    # 1. Introduction
    doc.add_heading("1. Introduction", level=1)

    doc.add_heading("1.1 Purpose", level=2)
    add_paragraphs(
        doc,
        "This SRS defines the functional and non-functional requirements for the "
        "Venue Booking System built for Central University. It is intended for "
        "developers, designers, supervisors, evaluators, and stakeholders who need "
        "a complete reference for implementation, testing, and deployment.",
    )

    doc.add_heading("1.2 Scope of the System", level=2)
    add_paragraphs(
        doc,
        "The system is a web-based SPA (React + Vite) that centralizes venue "
        "booking and allocation. It removes device-dependent spreadsheets or "
        "manual tracking by persisting bookings centrally so that availability is "
        "consistent across all devices and sessions in real time.",
    )

    doc.add_heading("1.3 Definitions, Acronyms, and Abbreviations", level=2)
    add_paragraphs(
        doc,
        "SRS: Software Requirements Specification\n"
        "RBAC: Role-Based Access Control\n"
        "Admin: User managing venues and approvals\n"
        "Venue: A bookable location (hall, classroom, auditorium)\n"
        "Timetable: A structured schedule showing venue usage",
    )

    doc.add_heading("1.4 References", level=2)
    add_paragraphs(
        doc,
        "IEEE 830 / IEEE 29148 SRS standards; Central University venue usage "
        "policies; Vercel deployment and routing docs; React Router documentation "
        "for SPA routing.",
    )

    # 2. Overall Description
    doc.add_heading("2. Overall Description", level=1)

    doc.add_heading("2.1 Product Perspective", level=2)
    add_paragraphs(
        doc,
        "Client–server web application. Frontend: React/Vite SPA with React Router "
        "and Tailwind UI. Backend/data layer: centralized storage (planned) to "
        "ensure consistent availability and booking states. Deployed on Vercel "
        "with SPA rewrites to support client-side routing.",
    )

    doc.add_heading("2.2 Product Functions", level=2)
    add_paragraphs(
        doc,
        "• Venue booking and scheduling\n"
        "• Availability checking and conflict prevention\n"
        "• Administrative approval workflows\n"
        "• Timetable visualization\n"
        "• User and role management",
    )

    doc.add_heading("2.3 User Classes and Characteristics", level=2)
    add_paragraphs(
        doc,
        "Student: requests venues for academic activities.\n"
        "Lecturer/Staff: books venues for lectures or events.\n"
        "Admin: manages venues, bookings, and approvals.\n"
        "Super Admin: full system control and configuration.",
    )

    doc.add_heading("2.4 Operating Environment", level=2)
    add_paragraphs(
        doc,
        "Modern browsers (Chrome, Firefox, Edge, Safari); desktop and mobile; "
        "internet-connected; deployed on Vercel.",
    )

    doc.add_heading("2.5 Constraints", level=2)
    add_paragraphs(
        doc,
        "Maximum booking duration limited to 6 hours.\n"
        "No double bookings per venue/time slot.\n"
        "Administrative approval may be required.\n"
        "SPA routing requires server rewrites (configured in vercel.json).",
    )

    doc.add_heading("2.6 Assumptions and Dependencies", level=2)
    add_paragraphs(
        doc,
        "Users have internet access.\n"
        "Venue lists and schedules are maintained by administrators.\n"
        "Authentication service and database are available.\n"
        "Vercel handles static hosting and SPA rewrites.",
    )

    # 3. System Requirements
    doc.add_heading("3. System Requirements", level=1)

    doc.add_heading("3.1 Functional Requirements", level=2)
    fr_groups = [
        (
            "Booking Management",
            [
                "FR-1: The system shall allow users to create venue booking requests.",
                "FR-2: The system shall prevent double booking of venues.",
                "FR-3: The system shall enforce a maximum booking duration of 6 hours.",
                "FR-4: The system shall display real-time booking status across all devices.",
                "FR-5: The system shall allow users to view booking history.",
            ],
        ),
        (
            "Venue Management",
            [
                "FR-6: The system shall allow admins to add, edit, and remove venues.",
                "FR-7: The system shall categorize venues by buildings (A–F).",
                "FR-8: The system shall define venue capacity and availability.",
            ],
        ),
        (
            "Timetable Management",
            [
                "FR-9: The system shall display bookings in a timetable format.",
                "FR-10: The system shall update timetables automatically after bookings.",
            ],
        ),
        (
            "Authentication and Authorization",
            [
                "FR-11: The system shall authenticate users before access.",
                "FR-12: The system shall enforce role-based access control.",
                "FR-13: The system shall restrict administrative features to admins only.",
            ],
        ),
        (
            "Admin Dashboard",
            [
                "FR-14: The system shall provide an admin dashboard.",
                "FR-15: The admin shall approve or reject booking requests.",
                "FR-16: The admin shall manage users and system configurations.",
            ],
        ),
    ]

    for title, items in fr_groups:
        doc.add_heading(title, level=3)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3.2 Non-Functional Requirements", level=2)
    nfr = [
        (
            "Performance",
            [
                "The system shall respond to user requests within acceptable time limits.",
                "The system shall support concurrent users.",
            ],
        ),
        (
            "Security",
            [
                "The system shall protect user authentication data.",
                "The system shall prevent unauthorized access.",
            ],
        ),
        (
            "Usability",
            [
                "The system shall be user-friendly and intuitive.",
                "The system shall be responsive on mobile devices.",
            ],
        ),
        (
            "Reliability",
            [
                "The system shall maintain data consistency.",
                "The system shall handle errors gracefully.",
            ],
        ),
        ("Scalability", ["The system shall support future expansion of venues and users."]),
    ]

    for title, items in nfr:
        doc.add_heading(title, level=3)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    # 4. External Interface Requirements
    doc.add_heading("4. External Interface Requirements", level=1)

    doc.add_heading("4.1 User Interface", level=2)
    add_paragraphs(
        doc,
        "Web-based graphical UI; admin dashboard for management tasks; responsive layout.",
    )

    doc.add_heading("4.2 Software Interfaces", level=2)
    add_paragraphs(
        doc,
        "Database management system; authentication services; deployment on Vercel with SPA rewrites.",
    )

    # 5. System Architecture Overview
    doc.add_heading("5. System Architecture Overview", level=1)
    add_paragraphs(
        doc,
        "Client–server architecture: React/Vite SPA frontend using React Router for routing and Tailwind for UI; "
        "backend services (planned) to manage business logic and persistence; deployment on Vercel with rewrites "
        "ensuring all paths serve index.html for client-side routing.",
    )

    # 6. Business Rules
    doc.add_heading("6. Business Rules", level=1)
    for item in [
        "A venue can only be booked once per time slot.",
        "Booking duration must not exceed 6 hours.",
        "Admin approval is required for final booking confirmation.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # 7. Error Handling
    doc.add_heading("7. Error Handling", level=1)
    add_paragraphs(
        doc,
        "The system shall display appropriate error messages for invalid inputs and log errors for administrative review. "
        "Client-side routing errors are mitigated through Vercel rewrites to ensure index.html is served on refresh.",
    )

    # 8. Future Enhancements
    doc.add_heading("8. Future Enhancements", level=1)
    for item in [
        "Email and SMS notifications.",
        "Calendar integration.",
        "Analytics and reporting features.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # 9. Appendix
    doc.add_heading("9. Appendix", level=1)

    doc.add_heading("9.1 Limitations", level=2)
    add_paragraphs(
        doc,
        "System depends on internet availability. Current deployment uses client-side routing with rewrites; backend "
        "persistence and auth integration assumed/planned.",
    )

    doc.add_heading("9.2 Conclusion", level=2)
    add_paragraphs(
        doc,
        "This SRS defines the requirements of the Venue Booking System and serves as a reference for development, "
        "evaluation, and future enhancements.",
    )

    doc.save(file_path)
    print(f"SRS written to: {file_path}")


if __name__ == "__main__":
    main()
